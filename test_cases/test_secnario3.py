import pytest
from openpyxl.workbook import Workbook
from pytest_selenium import driver
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from openpyxl import Workbook
from docx import Document

from base_page.inventory_page import inventory_page
from utilites.read_properties import Read_Config


class Testsecnario3:
    username = Read_Config.get_username()
    password = Read_Config.get_Password()
    URL = Read_Config.get_URL()


    def test_login(self, setup):
        self.driver = setup
        self.driver.get(self.URL)
        self.driver.implicitly_wait(5)
        self.ip = inventory_page(self.driver)
        self.ip.login_username(self.username)
        self.ip.login_password(self.password)
        self.ip.click_login_button()

        l = self.driver.find_elements(By.XPATH, "//div[@class='inventory_list']/div")
        doc = Document()
        doc.add_heading("Items in the Inventory", level=1)

        for i in l:
            image = self.driver.find_element(By.XPATH, "//div[@class = 'inventory_item_img']//a//img").get_attribute("src")
            # description = driver.find_element(By.XPATH,"//div[@class = 'inventory_item_description']").text
            label = self.driver.find_element(By.XPATH, "//div[@class = 'inventory_item_label']").text
            des = self.driver.find_element(By.XPATH, "//div[@class = 'inventory_item_desc']").text
            price = self.driver.find_element(By.XPATH, "//div[@class = 'pricebar']").text

            doc.add_paragraph(f" *{label}*")
            doc.add_paragraph(f"Description: {des}")
            doc.add_paragraph(f"Price: {price}")
            doc.add_paragraph(f"Image URL: {image}")
            doc.add_paragraph("-" * 50)
        doc.save("inventory.docx")

        self.ip.click_menu_tab()
        self.ip.click_logout_button()

        #self.driver.find_element(By.XPATH, "//a[@id = 'logout_sidebar_link']").click()
        button_text = self.ip.get_login_button_class()

        #button_text = self.driver.find_element(By.XPATH, "//input[@type='submit']").get_attribute("class")
        self.driver.get_screenshot_as_file("screenshot/screenshot.png")

        if "button" in button_text:
            print("As this page has a login button that means it Login Page")




